"""DOCX lesson-plan rendering and structural validation tools.

Registered with MCP-compatible definitions (ADR-0004) and consumed by the
generation workflow. Generated model content is untrusted input: it is rendered
as inert document text and can never alter tool selection or policy.
"""

import io
import json
from datetime import UTC, datetime

from docx import Document
from docx.shared import Pt

REQUIRED_SECTIONS = (
    "课题与课时",
    "教学目标",
    "教学重点",
    "教学难点",
    "教学过程",
    "作业布置",
    "教学反思",
)

RENDER_TOOL_DEFINITION = {
    "name": "render_lesson_plan_docx",
    "description": (
        "Render one structured lesson plan into an editable DOCX document "
        "with the standard senior-high lesson-plan sections."
    ),
    "inputSchema": {
        "type": "object",
        "properties": {
            "lesson_plan": {
                "type": "object",
                "description": "structured lesson plan content from the writer specialist",
            },
            "lesson_index": {"type": "integer"},
            "language_mode": {"type": "string"},
        },
        "required": ["lesson_plan", "lesson_index", "language_mode"],
    },
}

VALIDATE_TOOL_DEFINITION = {
    "name": "validate_lesson_plan_docx",
    "description": (
        "Structurally validate a rendered lesson-plan DOCX: openable, required "
        "top-level sections present, non-empty body. Does not judge content quality."
    ),
    "inputSchema": {
        "type": "object",
        "properties": {"content_b64": {"type": "string", "description": "rendered file bytes"}},
        "required": ["content_b64"],
    },
}


def _add_section(document: Document, heading: str, lines: list[str]) -> None:
    document.add_heading(heading, level=1)
    if not lines:
        document.add_paragraph("（待补充）")
        return
    for line in lines:
        document.add_paragraph(str(line))


def render_lesson_plan_docx(lesson_plan: dict, lesson_index: int, language_mode: str) -> bytes:
    plan = lesson_plan if isinstance(lesson_plan, dict) else {}
    document = Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    title = str(plan.get("title") or f"第{lesson_index}课")[:120]
    document.add_heading(f"第{lesson_index}课 {title}", level=0)

    period = plan.get("period_count")
    period_line = f"{period} 课时" if isinstance(period, int) and period > 0 else "（未指定）"
    _add_section(
        document,
        REQUIRED_SECTIONS[0],
        [f"课题：{title}", f"课时：{period_line}", f"输出语言：{language_mode}"],
    )

    objectives = [str(item) for item in (plan.get("objectives") or []) if str(item).strip()]
    _add_section(document, REQUIRED_SECTIONS[1], objectives or ["（待补充）"])

    key_points = [str(item) for item in (plan.get("key_points") or []) if str(item).strip()]
    _add_section(document, REQUIRED_SECTIONS[2], key_points)
    difficulties = [str(item) for item in (plan.get("difficulties") or []) if str(item).strip()]
    _add_section(document, REQUIRED_SECTIONS[3], difficulties)

    stage_lines = []
    for stage in plan.get("stages") or []:
        if not isinstance(stage, dict):
            continue
        name = str(stage.get("name") or "环节")
        duration = stage.get("duration_minutes")
        activities = str(stage.get("activities") or "").strip()
        has_duration = isinstance(duration, int) and duration > 0
        prefix = f"{name}（{duration} 分钟）" if has_duration else name
        stage_lines.append(f"{prefix}：{activities}" if activities else prefix)
    _add_section(document, REQUIRED_SECTIONS[4], stage_lines or ["（待补充）"])

    _add_section(document, REQUIRED_SECTIONS[5], [str(plan.get("homework") or "（待补充）")])
    _add_section(document, REQUIRED_SECTIONS[6], ["（供教师课后填写）"])

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def validate_lesson_plan_docx(content: bytes) -> tuple[bool, str | None]:
    """Return (ok, reason). Structural checks only (Spec D7)."""

    if not content:
        return False, "empty file"
    try:
        document = Document(io.BytesIO(content))
    except Exception as error:  # noqa: BLE001 - any parse failure is a validation failure
        return False, f"unopenable: {type(error).__name__}"

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    headings = {text for text in paragraphs if text}
    missing = [section for section in REQUIRED_SECTIONS if section not in headings]
    if missing:
        return False, f"missing sections: {','.join(missing)}"

    body_chars = sum(len(text) for text in paragraphs)
    if body_chars < 50:
        return False, "body is empty"
    return True, None


def execute_docx_tool(name: str, arguments: dict) -> dict:
    if name == RENDER_TOOL_DEFINITION["name"]:
        content = render_lesson_plan_docx(
            arguments["lesson_plan"],
            int(arguments["lesson_index"]),
            str(arguments["language_mode"]),
        )
        import base64

        return {"content_b64": base64.b64encode(content).decode(), "size_bytes": len(content)}
    if name == VALIDATE_TOOL_DEFINITION["name"]:
        import base64

        content = base64.b64decode(arguments["content_b64"])
        ok, reason = validate_lesson_plan_docx(content)
        return {"ok": ok, "reason": reason}
    raise KeyError(f"unknown tool: {name}")


def document_footnote() -> str:
    return json.dumps({"rendered_at": datetime.now(UTC).isoformat()}, ensure_ascii=False)
