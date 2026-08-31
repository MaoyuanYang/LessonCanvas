"""DOCX exercise/answer pair rendering and deterministic pair validation tools.

Registered with MCP-compatible definitions (ADR-0004) and consumed by the
exercise workflow. The renderer owns continuous item numbering across the
whole set (Spec D1); the validator independently re-checks both rendered files
so a broken draft can never reach ready state (Spec D7). Generated model
content is untrusted input: it is rendered as inert document text and can
never alter tool selection or policy.
"""

import io
import re
from datetime import UTC, datetime

from docx import Document
from docx.shared import Pt

EXERCISE_INSTRUCTIONS_SECTION = "练习说明"
ANSWER_SECTION = "答案"

DIFFICULTY_LABELS = {
    "foundation": "基础",
    "consolidation": "巩固",
    "advanced": "进阶",
}

EXERCISE_CATALOG = {
    "multiple_choice": "选择题",
    "fill_in_the_blank": "填空题",
    "short_answer": "简答题",
    "reading_comprehension": "阅读理解",
    "translation": "翻译",
    "written_expression": "书面表达",
}

# Item identity is the leading "N." at paragraph start; the content may span
# multiple lines (writing-task reference answers are naturally multi-line), so
# the capture crosses newlines with DOTALL instead of anchoring at line end.
_ITEM_LINE = re.compile(r"^(\d+)\.\s*(.*)", re.DOTALL)

RENDER_TOOL_DEFINITION = {
    "name": "render_lesson_exercises_docx",
    "description": (
        "Render one structured exercise draft into the paired editable DOCX "
        "files (student exercise set and reference answer set) with renderer-"
        "owned continuous numbering shared by both files."
    ),
    "inputSchema": {
        "type": "object",
        "properties": {
            "exercise_set": {
                "type": "object",
                "description": "structured exercise+answer draft from the writer specialist",
            },
            "lesson_index": {"type": "integer"},
            "language_mode": {"type": "string"},
            "difficulty": {
                "type": "string",
                "enum": ["foundation", "consolidation", "advanced"],
            },
        },
        "required": ["exercise_set", "lesson_index", "language_mode", "difficulty"],
    },
}

VALIDATE_TOOL_DEFINITION = {
    "name": "validate_exercise_pair",
    "description": (
        "Deterministically validate a rendered exercise/answer DOCX pair: both "
        "files openable, required sections present, numbering contiguous from 1 "
        "within configured bounds, answer numbers equal to exercise numbers "
        "exactly, and every answer entry non-empty. Does not judge correctness "
        "or language quality."
    ),
    "inputSchema": {
        "type": "object",
        "properties": {
            "exercise_content_b64": {"type": "string"},
            "answer_content_b64": {"type": "string"},
        },
        "required": ["exercise_content_b64", "answer_content_b64"],
    },
}


def _flatten_items(exercise_set: dict) -> list[dict]:
    items: list[dict] = []
    for category in exercise_set.get("categories") or []:
        if not isinstance(category, dict):
            continue
        for item in category.get("items") or []:
            if isinstance(item, dict):
                items.append(item)
    return items


def _category_count(exercise_set: dict) -> int:
    return sum(
        1
        for category in exercise_set.get("categories") or []
        if isinstance(category, dict) and category.get("items")
    )


def render_exercise_pair(
    exercise_set: dict, lesson_index: int, language_mode: str, difficulty: str
) -> tuple[bytes, bytes]:
    """Render the paired exercise and answer DOCX files.

    The renderer assigns continuous numbering 1..N across all categories in
    draft order and applies the same numbers to the answer entries, so pairing
    exists by construction; the pair validator still re-checks the files.
    """

    draft = exercise_set if isinstance(exercise_set, dict) else {}
    title = str(draft.get("title") or f"第{lesson_index}课")[:120]
    tier_label = DIFFICULTY_LABELS.get(difficulty, difficulty)

    exercise_doc = Document()
    style = exercise_doc.styles["Normal"]
    style.font.size = Pt(11)
    exercise_doc.add_heading(f"第{lesson_index}课 {title} 练习", level=0)

    instructions = str(draft.get("instructions") or "（待补充）")
    exercise_doc.add_heading(EXERCISE_INSTRUCTIONS_SECTION, level=1)
    exercise_doc.add_paragraph(instructions)
    exercise_doc.add_paragraph(f"难度档位：{tier_label}")
    exercise_doc.add_paragraph(f"输出语言：{language_mode}")

    answer_doc = Document()
    answer_style = answer_doc.styles["Normal"]
    answer_style.font.size = Pt(11)
    answer_doc.add_heading(f"第{lesson_index}课 {title} 答案", level=0)
    answer_doc.add_heading(ANSWER_SECTION, level=1)

    number = 0
    for category in draft.get("categories") or []:
        if not isinstance(category, dict):
            continue
        type_name = EXERCISE_CATALOG.get(str(category.get("type")), str(category.get("type")))
        name = str(category.get("name") or type_name)
        items = [item for item in (category.get("items") or []) if isinstance(item, dict)]
        if not items:
            continue
        exercise_doc.add_heading(name, level=1)
        passage = str(category.get("passage") or "").strip()
        if passage:
            exercise_doc.add_paragraph(passage)
        for item in items:
            number += 1
            exercise_doc.add_paragraph(f"{number}. {item.get('stem') or '（待补充）'}")
            for option in item.get("options") or []:
                exercise_doc.add_paragraph(f"    {option}")
            answer_text = str(item.get("answer") or "").strip()
            rationale = str(item.get("rationale") or "").strip()
            entry = answer_text
            if answer_text and rationale:
                entry = f"{answer_text}（{rationale}）"
            answer_doc.add_paragraph(f"{number}. {entry}")

    buffer = io.BytesIO()
    exercise_doc.save(buffer)
    exercise_content = buffer.getvalue()

    buffer = io.BytesIO()
    answer_doc.save(buffer)
    answer_content = buffer.getvalue()
    return exercise_content, answer_content


def _numbered_entries(content: bytes) -> list[tuple[int, str]] | None:
    """Open one DOCX and return its (number, text) entries in document order.

    Returns None when the file cannot be opened by the parser.
    """

    if not content:
        return None
    try:
        document = Document(io.BytesIO(content))
    except Exception:  # noqa: BLE001 - any parse failure is a validation failure
        return None
    entries: list[tuple[int, str]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = _ITEM_LINE.match(text)
        if match:
            entries.append((int(match.group(1)), match.group(2).strip()))
    return entries


def _heading_texts(content: bytes, level: int) -> list[str] | None:
    try:
        document = Document(io.BytesIO(content))
    except Exception:  # noqa: BLE001 - any parse failure is a validation failure
        return None
    style_name = "Heading 1" if level == 1 else "Title"
    return [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name == style_name and paragraph.text.strip()
    ]


def validate_exercise_pair(
    exercise_content: bytes,
    answer_content: bytes,
    min_items: int | None = None,
    max_items: int | None = None,
    min_categories: int | None = None,
    max_categories: int | None = None,
) -> tuple[bool, str | None, dict]:
    """Return (ok, reason, stats). Deterministic structural and pairing checks
    only (Spec D7): every numbered exercise has exactly one non-empty answer
    entry and no orphan answers exist."""

    from lessoncanvas.settings import get_settings

    settings = get_settings()
    min_items = settings.exercise_min_items_per_lesson if min_items is None else min_items
    max_items = settings.exercise_max_items_per_lesson if max_items is None else max_items
    min_categories = (
        settings.exercise_min_categories_per_lesson
        if min_categories is None
        else min_categories
    )
    max_categories = (
        settings.exercise_max_categories_per_lesson
        if max_categories is None
        else max_categories
    )

    if not exercise_content:
        return False, "empty file", {"item_count": 0, "category_count": 0}
    if not answer_content:
        return False, "empty file", {"item_count": 0, "category_count": 0}

    exercise_entries = _numbered_entries(exercise_content)
    answer_entries = _numbered_entries(answer_content)
    if exercise_entries is None or answer_entries is None:
        return False, "unopenable file", {"item_count": 0, "category_count": 0}

    exercise_headings = _heading_texts(exercise_content, 1)
    answer_headings = _heading_texts(answer_content, 1)
    if exercise_headings is None or answer_headings is None:
        return False, "unopenable file", {"item_count": 0, "category_count": 0}
    if EXERCISE_INSTRUCTIONS_SECTION not in exercise_headings:
        return False, "missing exercise sections", {
            "item_count": 0,
            "category_count": 0,
        }
    if ANSWER_SECTION not in answer_headings:
        return False, "missing answer section", {"item_count": 0, "category_count": 0}

    category_count = sum(
        1 for heading in exercise_headings if heading != EXERCISE_INSTRUCTIONS_SECTION
    )
    item_count = len(exercise_entries)
    stats = {"item_count": item_count, "category_count": category_count}

    if not (min_items <= item_count <= max_items):
        return False, f"item count {item_count} outside bounds [{min_items},{max_items}]", stats
    if not (min_categories <= category_count <= max_categories):
        return (
            False,
            f"category count {category_count} outside bounds [{min_categories},{max_categories}]",
            stats,
        )

    exercise_numbers = [number for number, _ in exercise_entries]
    if exercise_numbers != list(range(1, item_count + 1)):
        return False, "exercise numbering is not contiguous from 1", stats

    answer_map: dict[int, str] = {}
    for number, text in answer_entries:
        if number in answer_map:
            return False, f"duplicate answer entry {number}", stats
        answer_map[number] = text
    exercise_set = set(exercise_numbers)
    answer_set = set(answer_map)
    if exercise_set != answer_set:
        missing = sorted(exercise_set - answer_set)
        orphan = sorted(answer_set - exercise_set)
        detail = []
        if missing:
            detail.append(f"missing answers: {missing}")
        if orphan:
            detail.append(f"orphan answers: {orphan}")
        return False, "; ".join(detail), stats

    empty = sorted(number for number, text in answer_map.items() if not text)
    if empty:
        return False, f"empty answer entries: {empty}", stats
    return True, None, stats


def execute_exercise_tool(name: str, arguments: dict) -> dict:
    import base64

    if name == RENDER_TOOL_DEFINITION["name"]:
        exercise_content, answer_content = render_exercise_pair(
            arguments["exercise_set"],
            int(arguments["lesson_index"]),
            str(arguments["language_mode"]),
            str(arguments["difficulty"]),
        )
        return {
            "exercise_content_b64": base64.b64encode(exercise_content).decode(),
            "answer_content_b64": base64.b64encode(answer_content).decode(),
            "exercise_size_bytes": len(exercise_content),
            "answer_size_bytes": len(answer_content),
        }
    if name == VALIDATE_TOOL_DEFINITION["name"]:
        ok, reason, stats = validate_exercise_pair(
            base64.b64decode(arguments["exercise_content_b64"]),
            base64.b64decode(arguments["answer_content_b64"]),
        )
        return {"ok": ok, "reason": reason, **stats}
    raise KeyError(f"unknown tool: {name}")


def pair_stats_of(exercise_set: dict) -> dict:
    """Draft-level counts used for event payloads before validation re-checks."""

    return {
        "item_count": len(_flatten_items(exercise_set)),
        "category_count": _category_count(exercise_set),
    }


def document_footnote() -> str:
    import json

    return json.dumps({"rendered_at": datetime.now(UTC).isoformat()}, ensure_ascii=False)
